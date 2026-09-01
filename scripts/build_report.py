#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""把诊断数据 JSON 生成单个 Word 报告（.docx）。

用法：
    python scripts/build_report.py 诊断数据.json 诊断报告.docx

诊断数据 JSON 的字段见同目录 report_schema.json（可复制后填写）。
依赖：python-docx（本机已装 1.2.0）。生成的 docx 用 Word 直接打开即可。
"""
import json
import sys

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

# A/B/C 三级的配色（字体颜色）
LEVEL_COLOR = {
    "A": RGBColor(0xC0, 0x00, 0x00),   # 红：否决性
    "B": RGBColor(0xBF, 0x6A, 0x00),   # 橙：重要
    "C": RGBColor(0x40, 0x60, 0x40),   # 灰绿：一般
}
HEADER_FILL = "1F4E79"   # 表头底色（深蓝）


def set_cell_bg(cell, hex_color):
    """给单元格填底色。"""
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.makeelement(qn("w:shd"), {
        qn("w:val"): "clear", qn("w:color"): "auto", qn("w:fill"): hex_color})
    tc_pr.append(shd)


def set_repeat_header(row):
    """让表头行在跨页时重复。"""
    tr_pr = row._tr.get_or_add_trPr()
    tr_pr.append(tr_pr.makeelement(qn("w:tblHeader"),
                                   {qn("w:val"): "true"}))


def add_table(doc, columns, rows, level_col=None):
    """加一张带边框、表头底色的表。level_col 指定哪列是分级（用于着色整行首列）。"""
    table = doc.add_table(rows=1, cols=len(columns))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0]
    set_repeat_header(hdr)
    for i, name in enumerate(columns):
        cell = hdr.cells[i]
        set_cell_bg(cell, HEADER_FILL)
        p = cell.paragraphs[0]
        run = p.add_run(str(name))
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.size = Pt(9)
    for r in rows:
        cells = table.add_row().cells
        for i, val in enumerate(r):
            p = cells[i].paragraphs[0]
            run = p.add_run("" if val is None else str(val))
            run.font.size = Pt(9)
            # 分级列着色
            if level_col is not None and i == level_col:
                run.bold = True
                run.font.color.rgb = LEVEL_COLOR.get(str(val).strip().upper(),
                                                      RGBColor(0, 0, 0))
    return table


def build_cover(doc, meta, conclusion):
    """封面 + 结论摘要。"""
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("安全设施设计专篇评审诊断报告")
    r.bold = True
    r.font.size = Pt(20)
    doc.add_paragraph()

    info = [
        ("项目名称", meta.get("project_name", "")),
        ("建设单位", meta.get("build_unit", "")),
        ("设计单位", meta.get("design_unit", "")),
        ("项目性质", meta.get("project_type", "")),
        ("属地", meta.get("location", "")),
        ("编制主依据", meta.get("compile_basis", "")),
        ("本次证据等级", meta.get("evidence_level", "")),
        ("本次评审标尺", meta.get("review_scale", "")),
        ("否决项阈值", meta.get("veto_threshold", "")),
        ("报告日期", meta.get("report_date", "")),
    ]
    t = doc.add_table(rows=0, cols=2)
    t.style = "Table Grid"
    for k, v in info:
        cells = t.add_row().cells
        kr = cells[0].paragraphs[0].add_run(k)
        kr.bold = True
        kr.font.size = Pt(10)
        set_cell_bg(cells[0], "DEEAF6")
        cells[1].paragraphs[0].add_run(str(v)).font.size = Pt(10)

    doc.add_paragraph()
    doc.add_heading("通过性结论", level=1)
    vp = doc.add_paragraph()
    vr = vp.add_run("结论：" + conclusion.get("verdict", ""))
    vr.bold = True
    vr.font.size = Pt(13)
    vr.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)
    if conclusion.get("reason"):
        doc.add_paragraph("理由：" + conclusion["reason"])
    if conclusion.get("l3_statement"):
        pl = doc.add_paragraph()
        rl = pl.add_run("L3证据声明：" + conclusion["l3_statement"])
        rl.italic = True
    doc.add_page_break()


def main():
    if len(sys.argv) < 3:
        print("用法：python build_report.py 诊断数据.json 诊断报告.docx")
        sys.exit(1)
    src, out = sys.argv[1], sys.argv[2]
    with open(src, "r", encoding="utf-8") as f:
        data = json.load(f)

    doc = Document()
    # 中文正文字体
    normal = doc.styles["Normal"]
    normal.font.name = "宋体"
    normal.font.size = Pt(10.5)
    normal.element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")

    meta = data.get("meta", {})
    conclusion = data.get("conclusion", {})
    build_cover(doc, meta, conclusion)

    # 各分项表
    for sec in data.get("sections", []):
        doc.add_heading(sec.get("title", "分项"), level=1)
        if sec.get("note"):
            note = doc.add_paragraph(sec["note"])
            note.runs[0].italic = True
        add_table(doc, sec.get("columns", []), sec.get("rows", []))
        doc.add_paragraph()

    # A/B/C 缺陷清单
    defects = data.get("defects", [])
    if defects:
        doc.add_heading("缺陷清单（A/B/C 分级）", level=1)
        cols = ["ID", "级别", "专篇原文摘录", "缺陷描述", "依据/复算",
                "证据等级", "整改建议"]
        rows = [[d.get("id", ""), d.get("level", ""), d.get("excerpt", ""),
                 d.get("description", ""), d.get("basis", ""),
                 d.get("evidence_level", ""), d.get("suggestion", "")]
                for d in defects]
        add_table(doc, cols, rows, level_col=1)

    doc.save(out)
    print("已生成：%s（分项 %d 张表，缺陷 %d 条）"
          % (out, len(data.get("sections", [])), len(defects)))


if __name__ == "__main__":
    main()
