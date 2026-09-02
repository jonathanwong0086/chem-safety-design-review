#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""把一组分章 markdown 文件按索引顺序合并、排版成单个 Word 报告（.docx）。

用法：
    python scripts/merge_md_to_docx.py 索引文件.md 诊断报告.docx

其中“索引文件”就是文档组里的 00-索引.md：
  - 文件顶部用 --- 包起来的 key: value 块提供封面信息（项目名称、证据等级等）；
  - 正文里的有序/无序列表，每一项写一个分章 md 的文件名（相对索引文件所在目录），
    合并时按这个顺序依次拼接。

本脚本只依赖 python-docx（本机已装 1.2.0），自带一个轻量 markdown 解析，
不需要 pandoc / libreoffice。支持的 markdown 子集：
  # / ## / ### 标题、管道表格、有序与无序列表、> 引用、--- 分隔线、**加粗**。
"""
import os
import re
import sys

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

# A/B/C 三级的字体配色
LEVEL_COLOR = {
    "A": RGBColor(0xC0, 0x00, 0x00),   # 红：否决性
    "B": RGBColor(0xBF, 0x6A, 0x00),   # 橙：重要
    "C": RGBColor(0x40, 0x60, 0x40),   # 灰绿：一般
}
HEADER_FILL = "1F4E79"   # 表头底色（深蓝）
BODY_FONT = "宋体"

# 识别“分级/级别”列，用于给缺陷表着色
LEVEL_COL_PAT = re.compile(r"(分级|级别)")
BOLD_PAT = re.compile(r"\*\*(.+?)\*\*")


def set_cell_bg(cell, hex_color):
    """给单元格填底色。"""
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.makeelement(qn("w:shd"), {
        qn("w:val"): "clear", qn("w:color"): "auto", qn("w:fill"): hex_color})
    tc_pr.append(shd)


def set_repeat_header(row):
    """让表头行在跨页时重复出现。"""
    tr_pr = row._tr.get_or_add_trPr()
    tr_pr.append(tr_pr.makeelement(qn("w:tblHeader"), {qn("w:val"): "true"}))


def add_runs_with_bold(paragraph, text, base_size=Pt(10.5)):
    """把 **加粗** 语法转成实际加粗 run，其余为普通 run。"""
    pos = 0
    for m in BOLD_PAT.finditer(text):
        if m.start() > pos:
            r = paragraph.add_run(text[pos:m.start()])
            r.font.size = base_size
        r = paragraph.add_run(m.group(1))
        r.bold = True
        r.font.size = base_size
        pos = m.end()
    if pos < len(text):
        r = paragraph.add_run(text[pos:])
        r.font.size = base_size
    return paragraph


def add_md_table(doc, header, rows):
    """把一张 markdown 管道表格渲染成带边框、表头底色的 Word 表。
    如果表头里有“级别/分级”列，该列按 A/B/C 分色加粗。"""
    level_col = None
    for i, h in enumerate(header):
        if LEVEL_COL_PAT.search(h):
            level_col = i
            break
    table = doc.add_table(rows=1, cols=len(header))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0]
    set_repeat_header(hdr)
    for i, name in enumerate(header):
        cell = hdr.cells[i]
        set_cell_bg(cell, HEADER_FILL)
        p = cell.paragraphs[0]
        run = p.add_run(name)
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.size = Pt(9)
    for r in rows:
        cells = table.add_row().cells
        for i in range(len(header)):
            val = r[i] if i < len(r) else ""
            p = cells[i].paragraphs[0]
            if level_col is not None and i == level_col:
                run = p.add_run(val)
                run.font.size = Pt(9)
                run.bold = True
                run.font.color.rgb = LEVEL_COLOR.get(
                    val.strip().upper(), RGBColor(0, 0, 0))
            else:
                add_runs_with_bold(p, val, base_size=Pt(9))
    return table


def build_cover(doc, meta):
    """用索引 frontmatter 生成封面 + 信息表。"""
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run(meta.get("报告标题", "安全设施设计专篇评审诊断报告"))
    r.bold = True
    r.font.size = Pt(20)
    doc.add_paragraph()

    # 除“报告标题”外的字段都进信息表，按索引里出现的顺序
    info = [(k, v) for k, v in meta.items() if k != "报告标题"]
    if info:
        t = doc.add_table(rows=0, cols=2)
        t.style = "Table Grid"
        for k, v in info:
            cells = t.add_row().cells
            kr = cells[0].paragraphs[0].add_run(k)
            kr.bold = True
            kr.font.size = Pt(10)
            set_cell_bg(cells[0], "DEEAF6")
            cells[1].paragraphs[0].add_run(v).font.size = Pt(10)
    doc.add_page_break()


def parse_frontmatter(text):
    """解析文件顶部 --- 包起来的 key: value 块，返回 (meta字典, 去掉块后的正文)。"""
    meta = {}
    lines = text.splitlines()
    if lines and lines[0].strip() == "---":
        end = None
        for i in range(1, len(lines)):
            if lines[i].strip() == "---":
                end = i
                break
        if end is not None:
            for ln in lines[1:end]:
                if ":" in ln:
                    k, v = ln.split(":", 1)
                    meta[k.strip()] = v.strip()
            return meta, "\n".join(lines[end + 1:])
    return meta, text


def split_table_row(line):
    """拆一行管道表格，去掉首尾空管道。"""
    cells = line.strip().strip("|").split("|")
    return [c.strip() for c in cells]


def is_separator_row(line):
    """判断是不是表格的 |---|---| 分隔行。"""
    s = line.strip().strip("|")
    return bool(s) and set(s.replace("|", "").replace(":", "").strip()) <= {"-", " "}


def render_markdown(doc, md_text):
    """把一章 markdown 正文渲染进 doc（就地追加）。"""
    lines = md_text.splitlines()
    i = 0
    n = len(lines)
    while i < n:
        line = lines[i]
        stripped = line.strip()
        # 空行
        if not stripped:
            i += 1
            continue
        # 标题
        m = re.match(r"^(#{1,6})\s+(.*)$", stripped)
        if m:
            level = min(len(m.group(1)), 3)
            doc.add_heading(m.group(2).strip(), level=level)
            i += 1
            continue
        # 分隔线
        if stripped in ("---", "***", "___"):
            i += 1
            continue
        # 表格：当前行是管道行，且下一行是分隔行
        if stripped.startswith("|") and i + 1 < n and is_separator_row(lines[i + 1]):
            header = split_table_row(lines[i])
            rows = []
            i += 2
            while i < n and lines[i].strip().startswith("|"):
                rows.append(split_table_row(lines[i]))
                i += 1
            add_md_table(doc, header, rows)
            doc.add_paragraph()
            continue
        # 引用块
        if stripped.startswith(">"):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Pt(18)
            add_runs_with_bold(p, stripped.lstrip(">").strip())
            p.runs and setattr(p.runs[0].font, "italic", True)
            i += 1
            continue
        # 有序列表
        mo = re.match(r"^\d+[\.、]\s*(.*)$", stripped)
        if mo:
            p = doc.add_paragraph(style="List Number")
            add_runs_with_bold(p, mo.group(1))
            i += 1
            continue
        # 无序列表
        mu = re.match(r"^[-*+]\s+(.*)$", stripped)
        if mu:
            p = doc.add_paragraph(style="List Bullet")
            add_runs_with_bold(p, mu.group(1))
            i += 1
            continue
        # 普通段落
        p = doc.add_paragraph()
        add_runs_with_bold(p, stripped)
        i += 1


def extract_chapter_files(index_body):
    """从索引正文的列表项里抽出各分章 md 的文件名，保持出现顺序。"""
    files = []
    for line in index_body.splitlines():
        s = line.strip()
        m = re.match(r"^(?:\d+[\.、]|[-*+])\s+(.*)$", s)
        if not m:
            continue
        item = m.group(1).strip()
        # 支持 [显示名](文件.md) 或直接写 文件.md
        link = re.search(r"\(([^)]+\.md)\)", item)
        if link:
            files.append(link.group(1).strip())
        else:
            token = item.split()[0] if item.split() else ""
            if token.endswith(".md"):
                files.append(token)
    return files


def main():
    if len(sys.argv) < 3:
        print("用法：python merge_md_to_docx.py 索引文件.md 诊断报告.docx")
        sys.exit(1)
    index_path, out = sys.argv[1], sys.argv[2]
    base_dir = os.path.dirname(os.path.abspath(index_path))

    with open(index_path, "r", encoding="utf-8") as f:
        index_text = f.read()
    meta, index_body = parse_frontmatter(index_text)
    chapter_files = extract_chapter_files(index_body)

    doc = Document()
    # 中文正文字体
    normal = doc.styles["Normal"]
    normal.font.name = BODY_FONT
    normal.font.size = Pt(10.5)
    normal.element.rPr.rFonts.set(qn("w:eastAsia"), BODY_FONT)

    build_cover(doc, meta)

    missing = []
    merged = 0
    for idx, fn in enumerate(chapter_files):
        path = os.path.join(base_dir, fn)
        if not os.path.isfile(path):
            missing.append(fn)
            continue
        with open(path, "r", encoding="utf-8") as f:
            _, body = parse_frontmatter(f.read())
        if merged > 0:
            doc.add_page_break()
        render_markdown(doc, body)
        merged += 1

    doc.save(out)
    print("已生成：%s（合并 %d 章）" % (out, merged))
    if missing:
        print("警告：索引里列出但找不到的文件：%s" % "，".join(missing))


if __name__ == "__main__":
    main()
